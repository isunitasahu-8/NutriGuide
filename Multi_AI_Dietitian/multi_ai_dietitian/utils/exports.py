from __future__ import annotations

from typing import Dict, Any, List
import csv
import io
from datetime import datetime

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document


def _flatten_daily_meals(plan: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    daily = plan.get("daily_meals", {})
    for day_key, meals in daily.items():
        for meal_key, meal in meals.items():
            rows.append({
                "day": day_key,
                "meal": meal_key,
                "name": meal.get("name", ""),
                "calories": meal.get("calories", 0),
                "protein_g": meal.get("protein_g", 0),
                "carbs_g": meal.get("carbs_g", 0),
                "fats_g": meal.get("fats_g", 0),
            })
    return rows


def export_csv(plan: Dict[str, Any], analysis: Dict[str, Any]) -> bytes:
    buffer = io.StringIO()
    writer = csv.DictWriter(buffer, fieldnames=["day", "meal", "name", "calories", "protein_g", "carbs_g", "fats_g"]) 
    writer.writeheader()
    for row in _flatten_daily_meals(plan):
        writer.writerow(row)
    # Analysis summary rows
    writer.writerow({})
    writer.writerow({"day": "Summary", "meal": "Avg/day", "calories": analysis.get("daily", {}).get("summary", {}).get("avg_calories", 0)})
    return buffer.getvalue().encode("utf-8")


def export_pdf(plan: Dict[str, Any], analysis: Dict[str, Any]) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    def draw_header(canvas, title, y_pos, font_size=20):
        canvas.setFont("Helvetica-Bold", font_size)
        canvas.drawString(50, y_pos, title)
        return y_pos - 35
    
    def draw_subheader(canvas, title, y_pos, font_size=14):
        canvas.setFont("Helvetica-Bold", font_size)
        canvas.drawString(50, y_pos, title)
        return y_pos - 25
    
    def draw_section_header(canvas, title, y_pos, font_size=12):
        canvas.setFont("Helvetica-Bold", font_size)
        canvas.drawString(50, y_pos, title)
        return y_pos - 20
    
    def draw_text(canvas, text, y_pos, font_size=10):
        canvas.setFont("Helvetica", font_size)
        canvas.drawString(50, y_pos, text)
        return y_pos - 15
    
    def draw_highlighted_text(canvas, text, y_pos, font_size=11):
        canvas.setFont("Helvetica-Bold", font_size)
        canvas.drawString(50, y_pos, text)
        return y_pos - 18
    
    def draw_line(canvas, y_pos, thickness=1):
        canvas.setLineWidth(thickness)
        canvas.line(50, y_pos, width-50, y_pos)
        canvas.setLineWidth(1)
        return y_pos - 15
    
    def draw_double_line(canvas, y_pos):
        canvas.setLineWidth(2)
        canvas.line(50, y_pos, width-50, y_pos)
        canvas.setLineWidth(1)
        return y_pos - 20
    
    def new_page_if_needed(canvas, y_pos, margin=50):
        if y_pos < margin:
            canvas.showPage()
            return height - 50
        return y_pos
    
    # Page 1: Cover Page
    y = height - 120
    y = draw_header(c, "AI DIETITIAN PRO", y, 28)
    y = draw_text(c, "Personalized Meal Planning Guide", y, 16)
    y -= 30
    y = draw_text(c, "Generated on: " + datetime.now().strftime("%B %d, %Y at %I:%M %p"), y, 12)
    y -= 50
    
    # Table of Contents
    y = draw_subheader(c, "TABLE OF CONTENTS", y, 16)
    y = draw_line(c, y, 2)
    y = draw_highlighted_text(c, "1. Executive Summary", y)
    y = draw_highlighted_text(c, "2. Daily Meal Plans", y)
    y = draw_highlighted_text(c, "3. Nutritional Analysis", y)
    y = draw_highlighted_text(c, "4. Shopping List", y)
    y = draw_highlighted_text(c, "5. Safety Guidelines", y)
    y = draw_highlighted_text(c, "6. Recommendations", y)
    
    # Page 2: Executive Summary
    c.showPage()
    y = height - 50
    y = draw_header(c, "EXECUTIVE SUMMARY", y, 22)
    y = draw_double_line(c, y)
    
    # Daily nutrition summary
    y = draw_subheader(c, "Daily Nutritional Targets", y, 16)
    y = draw_line(c, y)
    y = draw_highlighted_text(c, f"Total Daily Calories: {plan.get('total_calories', 0):.0f} kcal", y)
    y = draw_highlighted_text(c, f"Total Daily Protein: {plan.get('total_protein', 0):.1f} g", y)
    y = draw_highlighted_text(c, f"Total Daily Carbohydrates: {plan.get('total_carbs', 0):.1f} g", y)
    y = draw_highlighted_text(c, f"Total Daily Fats: {plan.get('total_fats', 0):.1f} g", y)
    y -= 15
    
    # Analysis summary
    daily = analysis.get("daily", {}).get("summary", {})
    if daily:
        y = draw_subheader(c, "Nutritional Analysis", y, 16)
        y = draw_line(c, y)
        y = draw_highlighted_text(c, f"Average Daily Calories: {daily.get('avg_calories', 0):.0f} kcal", y)
        y = draw_highlighted_text(c, f"Average Daily Protein: {daily.get('avg_protein_g', 0):.1f} g", y)
        y = draw_highlighted_text(c, f"Average Daily Carbohydrates: {daily.get('avg_carbs_g', 0):.1f} g", y)
        y = draw_highlighted_text(c, f"Average Daily Fats: {daily.get('avg_fats_g', 0):.1f} g", y)
        y -= 15
    
    # Cost and sustainability
    cost_data = analysis.get("cost", {})
    sustainability_data = analysis.get("sustainability", {})
    if cost_data or sustainability_data:
        y = draw_subheader(c, "Cost & Environmental Impact", y, 16)
        y = draw_line(c, y)
        if cost_data:
            y = draw_highlighted_text(c, f"Estimated Daily Cost: ${cost_data.get('average_cost_per_day', 0):.2f}", y)
        if sustainability_data:
            y = draw_highlighted_text(c, f"Daily CO2 Emissions: {sustainability_data.get('average_kg_co2e_per_day', 0):.2f} kg", y)
        y -= 15
    
    # Page 3+: Daily Meal Plans
    c.showPage()
    y = height - 50
    y = draw_header(c, "DAILY MEAL PLANS", y, 22)
    y = draw_double_line(c, y)
    
    daily_meals = plan.get("daily_meals", {})
    for day_key, meals in daily_meals.items():
        day_name = day_key.replace('_', ' ').title()
        y = new_page_if_needed(c, y)
        y = draw_subheader(c, f"{day_name}", y, 18)
        y = draw_line(c, y, 2)
        
        for meal_type, meal in meals.items():
            meal_name = meal_type.replace('_', ' ').title()
            y = new_page_if_needed(c, y)
            y = draw_section_header(c, f"{meal_name}: {meal.get('name', '')}", y, 14)
            
            # Nutrition info
            y = draw_highlighted_text(c, f"Calories: {meal.get('calories', 0):.0f} kcal", y)
            y = draw_highlighted_text(c, f"Protein: {meal.get('protein_g', 0):.1f} g", y)
            y = draw_highlighted_text(c, f"Carbohydrates: {meal.get('carbs_g', 0):.1f} g", y)
            y = draw_highlighted_text(c, f"Fats: {meal.get('fats_g', 0):.1f} g", y)
            
            # Ingredients
            ingredients = meal.get("ingredients", [])
            if ingredients:
                y = draw_highlighted_text(c, "Ingredients:", y)
                for ing in ingredients:
                    amount = ing.get('amount', ing.get('grams', ''))
                    method = ing.get('method', '')
                    if amount and method:
                        y = draw_text(c, f"• {ing.get('name', '')} - {amount}g ({method})", y)
                    elif amount:
                        y = draw_text(c, f"• {ing.get('name', '')} - {amount}g", y)
                    else:
                        y = draw_text(c, f"• {ing.get('name', '')}", y)
                    y = new_page_if_needed(c, y)
            
            # Instructions
            instructions = meal.get("instructions", "")
            if instructions:
                y = draw_highlighted_text(c, "Instructions:", y)
                # Split long instructions into multiple lines
                words = instructions.split()
                line = ""
                for word in words:
                    if len(line + word) < 80:
                        line += word + " "
                    else:
                        y = draw_text(c, line.strip(), y)
                        y = new_page_if_needed(c, y)
                        line = word + " "
                if line:
                    y = draw_text(c, line.strip(), y)
            
            y -= 15
    
    # Shopping List Page
    c.showPage()
    y = height - 50
    y = draw_header(c, "SHOPPING LIST", y, 22)
    y = draw_double_line(c, y)
    
    # Generate shopping list
    shopping_categories = {
        "Proteins": [],
        "Grains & Carbs": [],
        "Vegetables": [],
        "Fruits": [],
        "Dairy": [],
        "Fats & Oils": [],
        "Nuts & Seeds": [],
        "Other": []
    }
    
    # Categorize ingredients
    for day_meals in daily_meals.values():
        for meal in day_meals.values():
            for ing in meal.get("ingredients", []):
                name = ing.get('name', '').lower()
                amount = ing.get('grams', ing.get('amount', ''))
                
                if amount:
                    item = f"{name.title()} ({amount}g)"
                else:
                    item = name.title()
                
                # Simple categorization
                if any(p in name for p in ["chicken", "salmon", "tofu", "lentils", "eggs", "greek yogurt", "fish", "beef", "pork"]):
                    if item not in shopping_categories["Proteins"]:
                        shopping_categories["Proteins"].append(item)
                elif any(g in name for g in ["brown rice", "quinoa", "oats", "bread", "pasta", "rice"]):
                    if item not in shopping_categories["Grains & Carbs"]:
                        shopping_categories["Grains & Carbs"].append(item)
                elif any(v in name for v in ["broccoli", "spinach", "tomato", "carrots", "bell pepper", "onion", "garlic"]):
                    if item not in shopping_categories["Vegetables"]:
                        shopping_categories["Vegetables"].append(item)
                elif any(f in name for f in ["banana", "apple", "berries", "orange", "grape"]):
                    if item not in shopping_categories["Fruits"]:
                        shopping_categories["Fruits"].append(item)
                elif any(d in name for d in ["milk", "cheese", "yogurt", "butter"]):
                    if item not in shopping_categories["Dairy"]:
                        shopping_categories["Dairy"].append(item)
                elif any(f in name for f in ["olive oil", "coconut oil", "avocado"]):
                    if item not in shopping_categories["Fats & Oils"]:
                        shopping_categories["Fats & Oils"].append(item)
                elif any(n in name for n in ["almonds", "walnuts", "chia seeds", "flax seeds"]):
                    if item not in shopping_categories["Nuts & Seeds"]:
                        shopping_categories["Nuts & Seeds"].append(item)
                else:
                    if item not in shopping_categories["Other"]:
                        shopping_categories["Other"].append(item)
    
    # Display shopping list
    for category, items in shopping_categories.items():
        if items:
            y = new_page_if_needed(c, y)
            y = draw_subheader(c, category, y, 16)
            y = draw_line(c, y)
            for item in items:
                y = draw_text(c, f"• {item}", y)
                y = new_page_if_needed(c, y)
            y -= 10
    
    # Recommendations Page
    c.showPage()
    y = height - 50
    y = draw_header(c, "RECOMMENDATIONS & GUIDELINES", y, 22)
    y = draw_double_line(c, y)
    
    y = draw_subheader(c, "General Guidelines", y, 16)
    y = draw_line(c, y)
    y = draw_text(c, "• Follow the meal plan as closely as possible for best results", y)
    y = draw_text(c, "• Stay hydrated by drinking 8-10 glasses of water daily", y)
    y = draw_text(c, "• Listen to your body and adjust portion sizes as needed", y)
    y = draw_text(c, "• Consult with a healthcare provider before making major dietary changes", y)
    y -= 15
    
    y = draw_subheader(c, "Storage Tips", y, 16)
    y = draw_line(c, y)
    y = draw_text(c, "• Store fresh ingredients in appropriate containers", y)
    y = draw_text(c, "• Freeze meal prep items for longer shelf life", y)
    y = draw_text(c, "• Keep track of expiration dates", y)
    y -= 15
    
    y = draw_subheader(c, "Contact Information", y, 16)
    y = draw_line(c, y)
    y = draw_text(c, "For questions or concerns about this meal plan, please consult", y)
    y = draw_text(c, "with a registered dietitian or healthcare provider.", y)
    
    c.save()
    return buffer.getvalue()


def export_docx(plan: Dict[str, Any], analysis: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading("AI Dietitian Pro - Meal Plan", 0)
    doc.add_paragraph(f"Daily calories: {plan.get('total_calories', 0)}")
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    for i, h in enumerate(["Day", "Meal", "Name", "Calories", "Protein(g)", "Carbs(g)"]):
        hdr[i].text = h
    for r in _flatten_daily_meals(plan)[:80]:
        row = table.add_row().cells
        row[0].text = str(r["day"]) 
        row[1].text = str(r["meal"]) 
        row[2].text = str(r["name"]) 
        row[3].text = str(int(r["calories"]))
        row[4].text = str(int(r["protein_g"]))
        row[5].text = str(int(r["carbs_g"]))
    doc.add_page_break()
    doc.add_heading("Analysis Summary", level=1)
    daily = analysis.get("daily", {}).get("summary", {})
    doc.add_paragraph(f"Avg calories/day: {daily.get('avg_calories', 0):.0f}")
    doc.add_paragraph(f"Avg cost/day: ${analysis.get('cost', {}).get('average_cost_per_day', 0):.2f}")
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


